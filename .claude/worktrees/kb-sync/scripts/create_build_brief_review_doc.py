from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("deliverables")
DOCX_PATH = OUT_DIR / "QuoteMate_Build_Brief_Review_2026-05-18.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(85, 85, 85)
LIGHT_FILL = "F2F4F7"


def set_run_font(run, name="Arial", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_paragraph_border_bottom(paragraph, color="D9E2F3", size="8", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    p.add_run(text)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    p.add_run(text)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def build_document():
    OUT_DIR.mkdir(exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Title", 23, BLACK, 0, 6),
        ("Heading 1", 16, BLUE, 12, 6),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for list_style in ("List Bullet", "List Number"):
        styles[list_style].font.name = "Arial"
        styles[list_style]._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        styles[list_style]._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        styles[list_style].font.size = Pt(11)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "QuoteMate"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(hp.runs[0], size=9, color=GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("Build brief review | 18 May 2026")
    set_run_font(run, size=9, color=GRAY)

    title = doc.add_paragraph(style="Title")
    title.add_run("QuoteMate Build Brief Review")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("Week of 18 May 2026 | Transcript fidelity check")
    set_run_font(subtitle_run, size=12, color=GRAY, italic=True)
    set_paragraph_border_bottom(subtitle)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(14)
    for label, value in (
        ("Prepared from", "communication transcript + original weekly build list"),
        ("Purpose", "check whether the detailed brief exactly follows the source material"),
    ):
        r1 = meta.add_run(f"{label}: ")
        set_run_font(r1, size=10.5, bold=True)
        r2 = meta.add_run(f"{value}\n")
        set_run_font(r2, size=10.5)

    add_heading(doc, "Verdict", 1)
    verdict = doc.add_paragraph()
    verdict.add_run("No. ").bold = True
    verdict.add_run(
        "The brief follows the intent of the call very well, but it is not an exact transcript-faithful build brief yet. "
        "It combines John's direct requests with sensible engineering interpretations and a few scope changes."
    )

    add_heading(doc, "Where It Does Not Match Exactly", 1)
    mismatches = [
        (
            "Calendar / urgency is deprioritized.",
            "John's original list explicitly includes a calendar function to create urgency and lock in the inspection or job after deposit. "
            "The brief turns this into a separate multi-week job and only mentions a possible countdown if time allows. That is a reprioritization, not an exact restatement."
        ),
        (
            "Tentative ideas are written as confirmed changes.",
            'Dropping "Best" and changing $199 to $99 were both discussed tentatively in the call. They belong under decisions to confirm with John, not under confirmed changes.'
        ),
        (
            "Mid-chat SMS / MMS adds a new assumption.",
            "John asked for product photos or options during the conversation. The brief changes that into a specific implementation path involving WhatsApp or SMS plus a link page. "
            "That may be sensible, but it is not something John explicitly decided in the transcript."
        ),
        (
            "The non-conversion workflow is narrowed.",
            "John asked for a workflow that triggers human follow-up for people who did not accept. The brief narrows this to status tracking plus a dashboard list for week one. "
            "That is a reasonable delivery slice, but it is not the full requested outcome."
        ),
        (
            "Several details are code-derived, not transcript-derived.",
            'Examples include the "oldest price book" fallback, the 28% markup fallback, the exact "price-checker" behavior, storage-bucket design, three sample renders, and the quote-cost breakdown. '
            "These may be good engineering findings, but they should be labelled as implementation notes from the codebase rather than direct communication requirements."
        ),
        (
            "One statement overstates the current app state.",
            "The brief says quote status is barely tracked at all. The code already records some accepted and paid states. The more accurate statement is that the funnel is not yet reliable enough to power a clean VA follow-up queue."
        ),
    ]
    for title_text, detail in mismatches:
        p = add_bullet(doc, "")
        p.runs[0].text = ""
        bold = p.add_run(title_text + " ")
        bold.bold = True
        p.add_run(detail)

    doc.add_page_break()
    add_heading(doc, "What Matches The Transcript Well", 1)
    matching_items = [
        "The two core product goals: pricing accuracy and the customer wow moment.",
        "A per-operator materials catalogue.",
        "Brand priming such as Clipsal.",
        "Range-level pricing such as 2000 series versus Iconic.",
        "A strong bill of materials per job type.",
        "Expansion of the standard job catalogue.",
        "Rendering the quoted product into the customer's own photo.",
        "Different pricing for customer-supplied versus tradie-supplied products.",
        "Rough range estimates when an exact quote is not possible.",
        "Cost per quote / transaction.",
        "Human follow-up for non-converters.",
        "Mid-process images or product options to the customer.",
    ]
    for item in matching_items:
        add_bullet(doc, item)

    add_heading(doc, "Edits Needed Before Approval", 1)
    recommended_edits = [
        'Rename WP8 from "Smaller confirmed changes" to "Items requiring John confirmation."',
        "Add a dedicated calendar / urgency work package, or clearly mark it as deferred pending John's approval.",
        'Reword WP7 as the "week-one implementation slice" of the setter request rather than the whole requested feature.',
        'Split the document into two layers: "Directly requested by John" and "Engineering interpretation / recommended implementation."',
        "Reframe the MMS section as an open technical decision, not as settled transcript fact.",
        'Replace "status is barely tracked" with wording closer to: "basic accepted / paid fields exist, but the follow-up funnel is not yet operationalized into a reliable VA queue."',
    ]
    for item in recommended_edits:
        add_numbered(doc, item)

    add_heading(doc, "Bottom Line", 1)
    final_para = doc.add_paragraph()
    final_para.add_run("The brief is strong and directionally right. ").bold = True
    final_para.add_run(
        "I would call it roughly 85-90% faithful to the substance of the conversation, but not exact enough to claim that it perfectly follows the transcript and this week's original list. "
        "The cleanest version would preserve John's asks first, then layer the engineering plan underneath them so nobody confuses the request with the implementation."
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_document()
