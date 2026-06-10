from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("docs") / "deliverables"
DOCX_PATH = OUT_DIR / "QuoteMate_Revised_Build_Brief_2026-05-18.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(85, 85, 85)
LIGHT_FILL = "F2F4F7"


def set_run_font(run, name="Arial", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_border_bottom(paragraph, color="D9E2F3", size="8", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_table_cell_margins(cell)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    p.add_run(text)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    p.add_run(text)
    return p


def add_labelled_paragraph(doc, label, text):
    p = doc.add_paragraph()
    r = p.add_run(label + " ")
    r.bold = True
    p.add_run(text)
    return p


def add_work_package(
    doc,
    number,
    title,
    intro,
    what_it_is,
    current_state,
    involved,
    watch_out,
    done_when,
):
    add_heading(doc, f"WP{number} - {title}", 2)
    add_labelled_paragraph(doc, "Introduction:", intro)
    add_labelled_paragraph(doc, "What it is:", what_it_is)
    add_labelled_paragraph(doc, "What exists now:", current_state)
    add_labelled_paragraph(doc, "What is involved:", "")
    for item in involved:
        add_bullet(doc, item)
    add_labelled_paragraph(doc, "Watch out:", watch_out)
    add_labelled_paragraph(doc, "Done when:", done_when)


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Title", 23, BLACK, 0, 6),
        ("Heading 1", 16, BLUE, 12, 6),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for list_style in ("List Bullet", "List Number"):
        styles[list_style].font.name = "Arial"
        styles[list_style]._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        styles[list_style]._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        styles[list_style].font.size = Pt(11)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "QuoteMate"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(hp.runs[0], size=9, color=GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("Detailed build brief | 18 May 2026")
    set_run_font(run, size=9, color=GRAY)

    title = doc.add_paragraph(style="Title")
    title.add_run("QuoteMate Build Brief")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("Week of 18 May 2026 | Detailed version aligned to John's task list")
    set_run_font(subtitle_run, size=12, color=GRAY, italic=True)
    set_paragraph_border_bottom(subtitle)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(14)
    for label, value in (
        ("Prepared from", "John's pasted weekly task list plus the Friday product call transcript"),
        ("Purpose", "state exactly what is tasked, explain what each item means, and show what must be built"),
    ):
        r1 = meta.add_run(f"{label}: ")
        set_run_font(r1, size=10.5, bold=True)
        r2 = meta.add_run(f"{value}\n")
        set_run_font(r2, size=10.5)

    add_heading(doc, "Read This First", 1)
    p = doc.add_paragraph()
    p.add_run("John's list is the main source of truth for this week. ").bold = True
    p.add_run(
        "The brief below keeps every requested item visible and turns the shorthand list into buildable work. "
        "Two themes run through the week: quotes must be dependable, and the customer experience must feel vivid enough to sell the job."
    )

    add_heading(doc, "Coverage Check", 1)
    coverage_rows = [
        ("Estimation consistency and accuracy", "WP1"),
        ("Materials pricing book for each operator", "WP2"),
        ("Prime products by brand, e.g. Clipsal", "WP2"),
        ("Price by range within brand, e.g. 2000 vs Iconic", "WP2"),
        ("Strong consistent bill of materials for each job type", "WP3"),
        ("Source accurate list of job types", "WP3"),
        ("AI render of the product in the customer's own photo", "WP4"),
        ("Different pricing when customer supplies vs tradie supplies", "WP5"),
        ("Calendar / urgency / lock in after deposit", "WP6"),
        ("Setter function for non-conversions / VA follow-up", "WP7"),
        ("Cost per transaction / quote", "WP8"),
        ("SMS / MMS to client in the middle of the process", "WP9"),
    ]
    coverage = doc.add_table(rows=1, cols=2)
    coverage.style = "Table Grid"
    set_table_width(coverage, [6480, 2880])
    for idx, text in enumerate(("John's requested change", "Where it is handled")):
        cell = coverage.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, LIGHT_FILL)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=10.5, bold=True)
    for req, wp in coverage_rows:
        cells = coverage.add_row().cells
        cells[0].text = req
        cells[1].text = wp
        for cell in cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.10
                for run in p.runs:
                    set_run_font(run, size=10.5)

    add_heading(doc, "Recommended Build Order", 1)
    order = [
        "Fix accuracy foundations first.",
        "Build the operator-owned product and materials catalogue.",
        "Lock jobs to repeatable bills of materials and expand the job catalogue.",
        "Use the catalogue to power exact-product images and supply-scenario pricing.",
        "Add the customer-control flows: calendar / urgency, follow-up, and mid-process product options.",
        "Add quote-cost reporting in parallel because it is important but does not block the rest.",
    ]
    for item in order:
        add_numbered(doc, item)

    add_heading(doc, "Detailed Work Packages", 1)

    add_work_package(
        doc,
        1,
        "Estimation consistency and accuracy",
        "This is the first item John listed and the foundation for everything else. If the same tradie request can produce different answers, the product loses trust immediately.",
        "Make the estimator deterministic wherever the business rules are known, and remove silent ways it can use the wrong price basis.",
        "QuoteMate already has shared assemblies, shared materials, RAG over similar past quotes, and a grounding validator. The call also raised a tenant-related pricing concern, so ownership and fallback behavior must be audited before trusting later tests.",
        [
            "Audit tenant ownership for price books and remove any blank, cross-tenant, or unsafe fallback path.",
            "Make estimator lookups explicitly trade-safe and tenant-safe before a quote can be drafted.",
            "Add regression tests around pricing ownership, repeated runs, and inspection routing when safe pricing is unavailable.",
            "Create a small accuracy test set for the week's target jobs so later changes can be measured against the same baseline.",
        ],
        "Do this before calling later outputs 'accurate'. New catalogues and renders are useful only after the price foundation is trustworthy.",
        "The same clean intake repeatedly returns the same grounded price behavior, and misconfigured pricing routes safely instead of silently producing a wrong quote.",
    )

    add_work_package(
        doc,
        2,
        "Per-operator materials pricing book, brand priming, and range pricing",
        "These are separate bullets in John's list, but they belong in one build because they all describe the same operator-owned product catalogue.",
        "Give each operator their own product list, including the brand they prefer and the exact series or range they sell.",
        "The codebase has shared materials and tenant material preferences by brand, but it does not yet provide a full operator-owned product catalogue with range/series, product photos, and operator-specific pricing.",
        [
            "Create a tenant-owned materials / products table rather than relying only on shared rows.",
            "Store category, brand, range / series, supplier, active flag, product photo, and operator-specific prices.",
            "Extend preference handling so an operator can prefer both brand and range, such as Clipsal 2000 or Clipsal Iconic.",
            "Update material lookup so tenant-owned products are available to the estimator and shown ahead of generic shared choices when appropriate.",
            "Keep the data model compatible with later rendering and mid-process product-choice flows.",
        ],
        "This is a keystone package. If the catalogue is poorly modelled now, WP4, WP5, and WP9 all become messy later.",
        "An operator can configure the products they actually use, the estimator can choose the correct brand and range, and the same catalogue can drive both price and imagery.",
    )

    add_work_package(
        doc,
        3,
        "Fixed bills of materials and accurate job-type sourcing",
        "John wants the same job to price the same way every time and also wants a broader list of standard jobs so more requests can be handled confidently.",
        "Define each standard job as a structured job record plus a fixed list of required materials and quantities.",
        "There is already a job / assembly catalogue, including shared and tenant custom assemblies. What is still missing is a structured bill-of-materials relationship that locks the parts used for each job.",
        [
            "Create a structured job-to-materials table with quantities, optionality, and applicability rules.",
            "Use that table when generating quote lines instead of letting the model freely decide the parts each time.",
            "Import John's researched standard job list only after validating the source data and mapping it to the correct trade.",
            "Add repeat-run tests for representative jobs so bill-of-materials and total price stability can be checked over time.",
        ],
        "Sourcing the job list and building the software are different jobs. Bad source data will only make wrong quotes happen more consistently.",
        "A standard job produces the same required materials and a stable quote across repeated runs, and new job types can be imported without destabilizing existing ones.",
    )

    add_work_package(
        doc,
        4,
        "AI rendering of the exact product in the customer's own photo",
        "This is the second thing John said will sell the product: the customer sees the actual quoted product on their own sink, wall, ceiling, or room.",
        "Use the customer's uploaded photo as the scene and the selected catalogue product as the visual reference.",
        "The app already generates Gemini previews from customer photos, but the product identity is still mostly text-driven rather than anchored to a real operator-approved catalogue photo.",
        [
            "Link quote line items back to the selected catalogue product.",
            "Pass the product photo together with the customer's scene photo to the image-generation step.",
            "Use the selected product consistently across the preview and any supporting sample images.",
            "Add checks so a quoted Clipsal Iconic product does not render as a generic or mismatched fitting.",
        ],
        "This depends on WP2. Without a real product catalogue photo, the model can only guess.",
        "A customer looking at a tap, toilet, downlight, or power point preview can recognize the same product that is being quoted.",
    )

    add_work_package(
        doc,
        5,
        "Different pricing when the customer supplies the product",
        "John specifically called out that some jobs should price differently when the customer supplies the item versus when the tradie supplies it.",
        "Support two supply modes for the same job: tradie-supplied and customer-supplied.",
        "Shared materials already carry some product metadata, but the weekly requirement needs a clear operator-level price difference between supply-and-install and install-only.",
        [
            "Add separate price fields for tradie-supplied and customer-supplied scenarios in the operator catalogue.",
            "Capture supply mode during intake or selection and carry it into the quote calculation.",
            "Change quote wording so the customer can see when they are supplying the item themselves.",
            "Make sure later image rendering still uses the selected product when the customer has chosen one.",
        ],
        "Customer-supplied does not mean price-free. Labour, risk, and any safety wording still need to be explicit.",
        "The same product can produce the correct supply-and-install price or install-only price, and the quote clearly explains which scenario was used.",
    )

    add_work_package(
        doc,
        6,
        "Calendar, urgency, and lock-in after deposit",
        "John wants the customer to feel urgency and to be able to lock in either the inspection or the job after paying.",
        "Create a flow that moves a customer from quote to held offer to booked work, with visible timing and confirmed commitment.",
        "The app already has payment, booking pages, scheduled-at data, and booking confirmation messages. It does not yet expose a full urgency / held-until experience or a richer tradie availability model.",
        [
            "Define the urgency rule: for example, how long the quote or inspection price is held.",
            "Show the hold / expiry clearly on the quote experience and related messages.",
            "Model the handoff after deposit so payment leads into a clear booked or reserved state.",
            "Extend availability handling only as far as needed for the selected week's booking flow, rather than building a full calendar product all at once.",
        ],
        "A countdown without a real business rule is only decoration. Decide what is actually being held and what happens when it expires.",
        "A customer can see the urgency, pay, and move into a clear reserved / booked state without the tradie manually patching the process together.",
    )

    add_work_package(
        doc,
        7,
        "Setter function for non-conversions",
        "John wants a VA to follow up with people who received a quote but did not accept it.",
        "Create a reliable follow-up queue for quotes that were sent but not converted.",
        "Some quote status fields already exist, including paid and accepted paths, but there is no dedicated operational queue that tells a VA who needs attention next.",
        [
            "Make the quote lifecycle reliable enough to distinguish sent, viewed, accepted, and paid.",
            "Add a needs-follow-up view filtered by age and non-conversion status.",
            "Show the customer's details, quote summary, last activity, and a direct path to contact them.",
            "Leave room for later automation, but make the human VA workflow useful first.",
        ],
        "The follow-up list is only as good as the underlying status events. Fix event reliability before trusting the queue.",
        "A VA can open the dashboard and immediately see which customers received quotes, did not accept, and should be contacted.",
    )

    add_work_package(
        doc,
        8,
        "Cost per transaction / quote",
        "John wants to understand what each quote costs to produce so the business can see whether the workflow is profitable.",
        "Roll up the variable cost of generating each quote into one stored figure with a component breakdown.",
        "The system performs AI, image, SMS, and sometimes voice work, but there is not yet one obvious per-quote profitability number exposed to the operator.",
        [
            "Record the usage units needed for AI, image generation, SMS, and voice.",
            "Apply the current provider rates to convert those units into dollars.",
            "Store the total and the component breakdown against the quote.",
            "Expose the cost in the dashboard so John can compare quote value against quote cost.",
        ],
        "Keep provider rates centralized so future price changes do not require hunting through several code paths.",
        "Each quote has a visible total production cost and a useful breakdown by cost driver.",
    )

    add_work_package(
        doc,
        9,
        "SMS / MMS to the client in the middle of the process",
        "John asked for the customer to receive product photos or options during the conversation, not only after the quote is already built.",
        "Show real operator-approved product options mid-conversation and record the customer's choice.",
        "The app already supports SMS flows, inbound MMS, photo-request links, and WhatsApp fallback behavior. What is still missing is an outbound product-choice experience tied to the operator's own catalogue.",
        [
            "Define the moments in the conversation where a product choice should be offered.",
            "Use the operator-owned catalogue from WP2 as the only source for the images and labels shown.",
            "Choose the supported outbound channel for the experience, such as image message or SMS with a choice page, based on what is reliable for the tenant's number setup.",
            "Save the customer's selected option so it drives both quote pricing and later rendering.",
        ],
        "Do not show generic products the tradie does not actually sell. Also do not treat this as finished until the selected option flows through to both price and preview.",
        "The customer can choose between real products in the conversation, and the chosen product is the same one used in the quote and preview.",
    )

    add_heading(doc, "Additional Call Items, Not In The Pasted Weekly List", 1)
    additions = [
        'Rough-range estimates when exact pricing is not possible.',
        'Possibly changing the inspection offer from $199 to $99.',
        'Possibly simplifying quote tiers from Good / Better / Best to Good / Better.',
        'Backup-AI resilience if the primary model is unavailable.',
    ]
    for item in additions:
        add_bullet(doc, item)
    add_labelled_paragraph(
        doc,
        "How to treat these:",
        "They were discussed in the call, but they are not part of the exact pasted weekly task list. Keep them visible as candidate follow-ons or decisions for John, not as silent replacements for the requested work above.",
    )

    doc.add_page_break()
    add_heading(doc, "Suggested Weekly Sequence", 1)
    focus = doc.add_table(rows=1, cols=3)
    focus.style = "Table Grid"
    set_table_width(focus, [1800, 2400, 5160])
    for idx, text in enumerate(("Phase", "Work", "Why")):
        cell = focus.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, LIGHT_FILL)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=10.5, bold=True)
    rows = [
        ("1", "WP1", "Protect the accuracy foundation before layering on more features."),
        ("2", "WP2 + WP3", "Build the catalogue and fixed job structure that the rest depends on."),
        ("3", "WP4 + WP5 + WP9", "Use the catalogue to improve customer-facing product choice and visualization."),
        ("4", "WP6 + WP7", "Strengthen the conversion and follow-up flow after the quote exists."),
        ("Parallel", "WP8", "Track cost alongside the rest because it is additive and useful immediately."),
    ]
    for row in rows:
        cells = focus.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.10
                for run in p.runs:
                    set_run_font(run, size=10.5)

    add_heading(doc, "Bottom Line", 1)
    final_para = doc.add_paragraph()
    final_para.add_run("All twelve items from John's pasted weekly list are now represented directly in the brief. ").bold = True
    final_para.add_run(
        "The updated structure keeps the requested work first, explains what each task actually means, and separates extra call ideas from the work John explicitly asked to focus on this week."
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_document()
